# -*- coding: utf-8 -*-
"""
BiblioQuest Report
Generador de reporte Word y PDF para la Ficha BiblioQuest.

Lee:
data/processed/biblioquest_protocol.json

Genera:
data/exports/biblioquest_report.docx
data/exports/biblioquest_report.pdf
"""

from __future__ import annotations

from pathlib import Path
import json
import re
from typing import Any, Dict, List


BASE_DIR = Path(__file__).resolve().parent.parent
DATA_DIR = BASE_DIR / "data"
PROCESSED_DIR = DATA_DIR / "processed"
EXPORTS_DIR = DATA_DIR / "exports"
ASSETS_DIR = BASE_DIR / "assets" / "images"

DEFAULT_JSON_PATH = PROCESSED_DIR / "biblioquest_protocol.json"
DEFAULT_DOCX_PATH = EXPORTS_DIR / "biblioquest_report.docx"
DEFAULT_PDF_PATH = EXPORTS_DIR / "biblioquest_report.pdf"

LOGO_CANDIDATES = [
    ASSETS_DIR / "logo_bibliomap.png",
    ASSETS_DIR / "logo_bibliointel.png",
    ASSETS_DIR / "logo_ucin.png",
]


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def ensure_dirs() -> None:
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)


def load_protocol(json_path: Path = DEFAULT_JSON_PATH) -> Dict[str, Any]:
    if not json_path.exists():
        raise FileNotFoundError(
            f"No se encontro el protocolo BiblioQuest en: {json_path}. "
            "Primero ejecuta: python modules\\biblioquest.py"
        )

    return json.loads(json_path.read_text(encoding="utf-8"))


def get_logo_path() -> Path | None:
    for path in LOGO_CANDIDATES:
        if path.exists():
            return path
    return None


def list_to_text(items: Any) -> str:
    if not items:
        return "No especificado."
    if isinstance(items, list):
        return "\n".join([f"- {clean_text(item)}" for item in items if clean_text(item)])
    return clean_text(items)


def get_nested(data: Dict[str, Any], keys: List[str], default: Any = "") -> Any:
    current = data
    for key in keys:
        if not isinstance(current, dict):
            return default
        current = current.get(key, default)
    return current


def generate_word_report(
    record: Dict[str, Any],
    docx_path: Path = DEFAULT_DOCX_PATH,
) -> Path:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_SECTION
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise ImportError(
            "Falta instalar python-docx. Ejecuta: pip install python-docx"
        ) from exc

    ensure_dirs()

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"

    logo_path = get_logo_path()
    if logo_path:
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(str(logo_path), width=Inches(1.8))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Ficha BiblioQuest")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("De la observacion al objetivo")
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph("")

    intro = doc.add_paragraph()
    intro.add_run("Modulo: ").bold = True
    intro.add_run("BiblioQuest")

    note = doc.add_paragraph()
    note.add_run("Nota metodologica: ").bold = True
    note.add_run(
        record.get(
            "methodological_note",
            "Basado en la metodologia HOMI: sinergia de los aportes intelectuales "
            "de Hurtado de Barrera, Ozturk, Kocaman y Kanbach, y Mirabal Gonzalez. "
            "El paradigma inteligentizador se incorpora como marco conceptual propio en desarrollo."
        )
    )

    doc.add_heading("1. Identificacion inicial", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    initial_rows = [
        ("Fecha", record.get("created_at", "")),
        ("Investigador", record.get("researcher_name", "")),
        ("Proyecto", record.get("project_title", "")),
        ("Tema tentativo", record.get("tentative_topic", "")),
        ("Area", record.get("research_area", "")),
    ]

    for label, value in initial_rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = clean_text(value) or "No especificado."

    doc.add_heading("2. Observacion inicial y situacion generadora", level=1)

    sections = [
        ("Situacion generadora", record.get("generative_situation", "")),
        ("Curiosidad cientifica", record.get("scientific_curiosity", "")),
        ("Proposito cognitivo", record.get("cognitive_purpose", "")),
        ("Proyeccion transformadora", record.get("transformative_projection", "")),
    ]

    for heading, text in sections:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(clean_text(text) or "No especificado.")

    doc.add_heading("3. Componentes metodologicos segun Hurtado", level=1)

    hurtado_table = doc.add_table(rows=0, cols=2)
    hurtado_table.style = "Table Grid"

    hurtado_rows = [
        ("Evento de estudio", record.get("study_event", "")),
        ("Unidad de analisis", record.get("unit_of_analysis", "")),
        ("Contexto", record.get("context", "")),
        ("Temporalidad", record.get("temporality", "")),
    ]

    for label, value in hurtado_rows:
        row = hurtado_table.add_row().cells
        row[0].text = label
        row[1].text = clean_text(value) or "No especificado."

    doc.add_heading("4. Objetivo general de investigacion", level=1)
    doc.add_paragraph(clean_text(record.get("general_objective", "")) or "No especificado.")

    diagnostics = record.get("diagnostics", {})
    objective_validation = diagnostics.get("objective_validation", {})
    detected_level = objective_validation.get("detected_level", {})

    doc.add_heading("4.1 Validacion del objetivo", level=2)

    validation_table = doc.add_table(rows=0, cols=2)
    validation_table.style = "Table Grid"
    validation_rows = [
        ("Estado", objective_validation.get("status", "No evaluado")),
        ("Puntaje", f"{objective_validation.get('score', 'No evaluado')}/100"),
        ("Verbo detectado", detected_level.get("detected_verb", "No detectado")),
        ("Nivel segun Hurtado", detected_level.get("level", "No detectado")),
        ("Tipo de investigacion sugerido", ", ".join(detected_level.get("suggested_research_types", []))),
    ]

    for label, value in validation_rows:
        row = validation_table.add_row().cells
        row[0].text = label
        row[1].text = clean_text(value) or "No especificado."

    recommendations = objective_validation.get("recommendations", [])
    doc.add_heading("4.2 Recomendaciones sobre el objetivo", level=2)
    if recommendations:
        for item in recommendations:
            doc.add_paragraph(clean_text(item), style="List Bullet")
    else:
        doc.add_paragraph("No se registran recomendaciones criticas para el objetivo general.")

    doc.add_heading("5. Objetivos especificos", level=1)
    specific_objectives = record.get("specific_objectives", [])
    if specific_objectives:
        for item in specific_objectives:
            doc.add_paragraph(clean_text(item), style="List Number")
    else:
        doc.add_paragraph("No especificado.")

    doc.add_heading("6. Preguntas de investigacion", level=1)
    doc.add_heading("Pregunta principal", level=2)
    doc.add_paragraph(clean_text(record.get("main_question", "")) or "No especificado.")

    doc.add_heading("Preguntas especificas", level=2)
    specific_questions = record.get("specific_questions", [])
    if specific_questions:
        for item in specific_questions:
            doc.add_paragraph(clean_text(item), style="List Bullet")
    else:
        doc.add_paragraph("No especificado.")

    doc.add_heading("7. Resultados esperados", level=1)
    doc.add_paragraph(clean_text(record.get("expected_results", "")) or "No especificado.")

    doc.add_heading("8. Justificacion del uso de bibliometria", level=1)
    doc.add_paragraph(clean_text(record.get("bibliometric_justification", "")) or "No especificado.")

    doc.add_heading("9. Ajuste objetivo-pregunta-datos-analisis-interpretacion", level=1)
    doc.add_paragraph(clean_text(record.get("objective_question_data_analysis_fit", "")) or "No especificado.")

    ozturk_fit = diagnostics.get("ozturk_fit", {})
    doc.add_heading("9.1 Criterio de ajuste bibliometrico", level=2)

    fit_table = doc.add_table(rows=0, cols=2)
    fit_table.style = "Table Grid"
    fit_rows = [
        ("Estado", ozturk_fit.get("status", "No evaluado")),
        ("Puntaje", f"{ozturk_fit.get('score', 'No evaluado')}/100"),
    ]

    for label, value in fit_rows:
        row = fit_table.add_row().cells
        row[0].text = label
        row[1].text = clean_text(value) or "No especificado."

    doc.add_heading("10. Lectura I2E segun Mirabal", level=1)
    mirabal = diagnostics.get("mirabal_i2e_a3d_s3", {})

    for label in ["intorno", "entorno", "extorno"]:
        doc.add_heading(label.capitalize(), level=2)
        doc.add_paragraph(clean_text(mirabal.get(label, "")) or "No especificado.")

    doc.add_heading("11. A3D y 3S", level=1)

    doc.add_heading("A3D", level=2)
    a3d = mirabal.get("a3d", {})
    for key, value in a3d.items():
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(key.capitalize() + ": ").bold = True
        p.add_run(clean_text(value) or "No especificado.")

    doc.add_heading("3S", level=2)
    s3 = mirabal.get("s3", {})
    for key, value in s3.items():
        label = "Sostenimiento" if key == "sostenimiento" else key.capitalize()
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(label + ": ").bold = True
        p.add_run(clean_text(value) or "No especificado.")

    doc.add_heading("12. Novedad potencial e incertidumbre", level=1)
    novelty = diagnostics.get("novelty_uncertainty", {})

    novelty_table = doc.add_table(rows=0, cols=2)
    novelty_table.style = "Table Grid"
    novelty_rows = [
        ("Novedad potencial", novelty.get("novelty_level", "No evaluada")),
        ("Puntaje de novedad", f"{novelty.get('novelty_score', 'No evaluado')}/100"),
        ("Incertidumbre", novelty.get("uncertainty_level", "No evaluada")),
        ("Puntaje de incertidumbre", f"{novelty.get('uncertainty_score', 'No evaluado')}/100"),
        ("Senales detectadas", ", ".join(novelty.get("detected_signals", []))),
    ]

    for label, value in novelty_rows:
        row = novelty_table.add_row().cells
        row[0].text = label
        row[1].text = clean_text(value) or "No especificado."

    doc.add_paragraph(clean_text(novelty.get("warning", "")))

    doc.add_heading("13. Operacionalizacion preliminar de objetivos", level=1)
    operationalization = diagnostics.get("objective_operationalization", [])

    op_table = doc.add_table(rows=1, cols=6)
    op_table.style = "Table Grid"
    headers = [
        "Nº",
        "Objetivo especifico",
        "Estimacion",
        "Vision",
        "Visualizacion",
        "Evidencia posible",
    ]

    for idx, header in enumerate(headers):
        op_table.rows[0].cells[idx].text = header

    if operationalization:
        for row_data in operationalization:
            row = op_table.add_row().cells
            row[0].text = clean_text(row_data.get("objective_number", ""))
            row[1].text = clean_text(row_data.get("specific_objective", ""))
            row[2].text = clean_text(row_data.get("estimated_achievement", ""))
            row[3].text = clean_text(row_data.get("vision", ""))
            row[4].text = clean_text(row_data.get("visualization", ""))
            row[5].text = clean_text(row_data.get("possible_evidence", ""))

    doc.add_heading("14. Comparacion entre proposito y objetivo general", level=1)

    compare_table = doc.add_table(rows=0, cols=2)
    compare_table.style = "Table Grid"
    compare_rows = [
        ("Proposito de investigacion", record.get("cognitive_purpose", "")),
        ("Objetivo general de investigacion", record.get("general_objective", "")),
        (
            "Lectura BiblioQuest",
            "El proposito expresa la intencion amplia y el horizonte de sentido de la investigacion. "
            "El objetivo general expresa el logro cognitivo central, delimitado y evaluable que orienta "
            "el metodo, los datos, el analisis y la interpretacion.",
        ),
    ]

    for label, value in compare_rows:
        row = compare_table.add_row().cells
        row[0].text = label
        row[1].text = clean_text(value) or "No especificado."

    doc.add_heading("15. Recomendacion para BiblioSearch", level=1)
    recommendations_search = [
        "Usar el objetivo general y la pregunta principal como base para construir terminos nucleares.",
        "Derivar palabras clave desde el evento de estudio, la unidad de analisis, el contexto y la temporalidad.",
        "Traducir terminos principales al ingles cuando la fuente bibliografica lo requiera.",
        "Construir al menos tres versiones de busqueda: amplia, intermedia y precisa.",
        "Registrar cada version de busqueda para mantener trazabilidad.",
    ]

    for item in recommendations_search:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("16. Dictamen BiblioQuest", level=1)
    doc.add_paragraph(
        f"Estado del objetivo: {objective_validation.get('status', 'No evaluado')}."
    )
    doc.add_paragraph(
        f"Estado del ajuste bibliometrico: {ozturk_fit.get('status', 'No evaluado')}."
    )
    doc.add_paragraph(
        "Recomendacion general: avanzar a BiblioSearch solo cuando objetivo, preguntas, "
        "resultados esperados y justificacion bibliometrica esten suficientemente alineados."
    )

    doc.add_heading("17. Referencias metodologicas preliminares", level=1)
    references = [
        "Hurtado de Barrera, J. (2014). Como formular objetivos de investigacion: Un acercamiento desde la investigacion holistica. Quiron Ediciones.",
        "Ozturk, O., Kocaman, R., & Kanbach, D. K. (2024). How to design bibliometric research: An overview and a framework proposal. Review of Managerial Science, 18(11), 3333-3361. https://doi.org/10.1007/s11846-024-00738-0",
        "Mirabal Gonzalez, J. F. (s. f.). El arbol para la innovacion: Transformando la empresa a la economia digital. OEM.",
    ]

    for ref in references:
        doc.add_paragraph(ref)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "BiblioQuest - BiblioIntel. De la observacion al objetivo."

    doc.save(str(docx_path))
    return docx_path


def split_text_for_pdf(text: str, max_len: int = 95) -> List[str]:
    text = clean_text(text)
    if not text:
        return ["No especificado."]

    words = text.split()
    lines = []
    current = ""

    for word in words:
        if len(current) + len(word) + 1 <= max_len:
            current = f"{current} {word}".strip()
        else:
            lines.append(current)
            current = word

    if current:
        lines.append(current)

    return lines


def generate_pdf_report(
    record: Dict[str, Any],
    pdf_path: Path = DEFAULT_PDF_PATH,
) -> Path:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
            PageBreak,
            Image,
        )
    except ImportError as exc:
        raise ImportError(
            "Falta instalar reportlab. Ejecuta: pip install reportlab"
        ) from exc

    ensure_dirs()

    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="Small",
            parent=styles["Normal"],
            fontSize=7,
            leading=9,
        )
    )
    styles.add(
        ParagraphStyle(
            name="NormalJustified",
            parent=styles["Normal"],
            fontSize=8.5,
            leading=11,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SectionTitle",
            parent=styles["Heading1"],
            fontSize=12,
            leading=14,
            spaceBefore=8,
            spaceAfter=5,
        )
    )

    story = []

    logo_path = get_logo_path()
    if logo_path:
        story.append(Image(str(logo_path), width=1.7 * inch, height=0.75 * inch))
        story.append(Spacer(1, 8))

    story.append(Paragraph("<b>Ficha BiblioQuest</b>", styles["Title"]))
    story.append(Paragraph("<i>De la observacion al objetivo</i>", styles["Heading2"]))
    story.append(Spacer(1, 8))

    story.append(
        Paragraph(
            "<b>Nota metodologica:</b> "
            + clean_text(record.get("methodological_note", "")),
            styles["NormalJustified"],
        )
    )
    story.append(Spacer(1, 10))

    def add_section(title: str, body: str) -> None:
        story.append(Paragraph(title, styles["SectionTitle"]))
        story.append(Paragraph(clean_text(body) or "No especificado.", styles["NormalJustified"]))
        story.append(Spacer(1, 5))

    add_section("1. Tema tentativo", record.get("tentative_topic", ""))
    add_section("2. Situacion generadora", record.get("generative_situation", ""))
    add_section("3. Curiosidad cientifica", record.get("scientific_curiosity", ""))
    add_section("4. Proposito cognitivo", record.get("cognitive_purpose", ""))
    add_section("5. Proyeccion transformadora", record.get("transformative_projection", ""))

    story.append(Paragraph("6. Componentes metodologicos segun Hurtado", styles["SectionTitle"]))

    component_data = [
        ["Componente", "Contenido"],
        ["Evento de estudio", clean_text(record.get("study_event", ""))],
        ["Unidad de analisis", clean_text(record.get("unit_of_analysis", ""))],
        ["Contexto", clean_text(record.get("context", ""))],
        ["Temporalidad", clean_text(record.get("temporality", ""))],
    ]

    component_table = Table(component_data, colWidths=[1.7 * inch, 5.2 * inch])
    component_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(component_table)
    story.append(Spacer(1, 8))

    add_section("7. Objetivo general de investigacion", record.get("general_objective", ""))

    diagnostics = record.get("diagnostics", {})
    objective_validation = diagnostics.get("objective_validation", {})
    detected_level = objective_validation.get("detected_level", {})
    ozturk_fit = diagnostics.get("ozturk_fit", {})
    novelty = diagnostics.get("novelty_uncertainty", {})
    mirabal = diagnostics.get("mirabal_i2e_a3d_s3", {})

    validation_data = [
        ["Criterio", "Resultado"],
        ["Estado objetivo", clean_text(objective_validation.get("status", "No evaluado"))],
        ["Puntaje objetivo", f"{objective_validation.get('score', 'No evaluado')}/100"],
        ["Verbo detectado", clean_text(detected_level.get("detected_verb", "No detectado"))],
        ["Nivel Hurtado", clean_text(detected_level.get("level", "No detectado"))],
        ["Ajuste bibliometrico", clean_text(ozturk_fit.get("status", "No evaluado"))],
        ["Puntaje ajuste", f"{ozturk_fit.get('score', 'No evaluado')}/100"],
        ["Novedad potencial", clean_text(novelty.get("novelty_level", "No evaluada"))],
        ["Incertidumbre", clean_text(novelty.get("uncertainty_level", "No evaluada"))],
    ]

    validation_table = Table(validation_data, colWidths=[2.0 * inch, 4.9 * inch])
    validation_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(Paragraph("8. Diagnostico BiblioQuest", styles["SectionTitle"]))
    story.append(validation_table)
    story.append(Spacer(1, 8))

    story.append(Paragraph("9. Objetivos especificos", styles["SectionTitle"]))
    for item in record.get("specific_objectives", []):
        story.append(Paragraph("- " + clean_text(item), styles["NormalJustified"]))
    story.append(Spacer(1, 6))

    add_section("10. Pregunta principal", record.get("main_question", ""))

    story.append(Paragraph("11. Preguntas especificas", styles["SectionTitle"]))
    for item in record.get("specific_questions", []):
        story.append(Paragraph("- " + clean_text(item), styles["NormalJustified"]))
    story.append(Spacer(1, 6))

    add_section("12. Resultados esperados", record.get("expected_results", ""))
    add_section("13. Justificacion del uso de bibliometria", record.get("bibliometric_justification", ""))
    add_section(
        "14. Ajuste objetivo-pregunta-datos-analisis-interpretacion",
        record.get("objective_question_data_analysis_fit", ""),
    )

    story.append(PageBreak())

    story.append(Paragraph("15. Lectura I2E segun Mirabal", styles["SectionTitle"]))
    add_section("Intorno", mirabal.get("intorno", ""))
    add_section("Entorno", mirabal.get("entorno", ""))
    add_section("Extorno", mirabal.get("extorno", ""))

    story.append(Paragraph("16. A3D y 3S", styles["SectionTitle"]))
    a3d = mirabal.get("a3d", {})
    s3 = mirabal.get("s3", {})

    a3d_s3_data = [["Dimension", "Contenido"]]
    for key, value in a3d.items():
        a3d_s3_data.append([key.capitalize(), clean_text(value)])
    for key, value in s3.items():
        label = "Sostenimiento" if key == "sostenimiento" else key.capitalize()
        a3d_s3_data.append([label, clean_text(value)])

    a3d_s3_table = Table(a3d_s3_data, colWidths=[1.6 * inch, 5.3 * inch])
    a3d_s3_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(a3d_s3_table)
    story.append(Spacer(1, 8))

    story.append(Paragraph("17. Operacionalizacion preliminar de objetivos", styles["SectionTitle"]))

    op = diagnostics.get("objective_operationalization", [])
    op_data = [["Nº", "Objetivo", "Estimacion", "Vision", "Visualizacion"]]

    for row in op:
        op_data.append(
            [
                clean_text(row.get("objective_number", "")),
                clean_text(row.get("specific_objective", "")),
                clean_text(row.get("estimated_achievement", "")),
                clean_text(row.get("vision", "")),
                clean_text(row.get("visualization", "")),
            ]
        )

    op_table = Table(
        op_data,
        colWidths=[0.35 * inch, 2.6 * inch, 1.5 * inch, 1.5 * inch, 1.0 * inch],
    )
    op_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 6),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(op_table)
    story.append(Spacer(1, 8))

    add_section(
        "18. Advertencia sobre novedad cientifica",
        novelty.get(
            "warning",
            "La novedad cientifica debe validarse mediante busqueda sistematica, lectura critica y contraste experto.",
        ),
    )

    story.append(Paragraph("19. Recomendacion para BiblioSearch", styles["SectionTitle"]))
    search_recs = [
        "Usar el objetivo general y la pregunta principal como base para construir terminos nucleares.",
        "Derivar palabras clave desde el evento de estudio, unidad de analisis, contexto y temporalidad.",
        "Traducir terminos principales al ingles cuando la fuente bibliografica lo requiera.",
        "Construir al menos tres versiones de busqueda: amplia, intermedia y precisa.",
        "Registrar cada version de busqueda para mantener trazabilidad.",
    ]
    for item in search_recs:
        story.append(Paragraph("- " + item, styles["NormalJustified"]))

    story.append(Spacer(1, 10))
    story.append(Paragraph("20. Referencias metodologicas preliminares", styles["SectionTitle"]))
    refs = [
        "Hurtado de Barrera, J. (2014). Como formular objetivos de investigacion: Un acercamiento desde la investigacion holistica. Quiron Ediciones.",
        "Ozturk, O., Kocaman, R., & Kanbach, D. K. (2024). How to design bibliometric research: An overview and a framework proposal. Review of Managerial Science, 18(11), 3333-3361. https://doi.org/10.1007/s11846-024-00738-0",
        "Mirabal Gonzalez, J. F. (s. f.). El arbol para la innovacion: Transformando la empresa a la economia digital. OEM.",
    ]

    for ref in refs:
        story.append(Paragraph(ref, styles["Small"]))

    def add_page_number(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.drawCentredString(
            A4[0] / 2.0,
            0.35 * inch,
            f"BiblioQuest - BiblioIntel | Pagina {doc_obj.page}",
        )
        canvas.restoreState()

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    return pdf_path


def generate_biblioquest_reports(
    json_path: Path = DEFAULT_JSON_PATH,
    docx_path: Path = DEFAULT_DOCX_PATH,
    pdf_path: Path = DEFAULT_PDF_PATH,
) -> Dict[str, str]:
    record = load_protocol(json_path)

    docx_file = generate_word_report(record, docx_path)
    pdf_file = generate_pdf_report(record, pdf_path)

    return {
        "docx_path": str(docx_file),
        "pdf_path": str(pdf_file),
    }


if __name__ == "__main__":
    result = generate_biblioquest_reports()
    print("Reportes BiblioQuest generados correctamente.")
    print(f"Word generado en: {result['docx_path']}")
    print(f"PDF generado en: {result['pdf_path']}")